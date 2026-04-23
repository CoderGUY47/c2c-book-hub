from docx import Document
from docx.shared import Pt
import sys
import os

doc = Document()

# Add a Title
doc.add_heading('Book-Shop: Complete Pre-Defence Report', 0)

markdown_path = r'C:\Users\CODERGUY\.gemini\antigravity\brain\4a5ef817-975d-4478-bddb-9f329597b661\final_pre_defence_draft.md'

if not os.path.exists(markdown_path):
    print("Markdown file not found.")
    sys.exit(1)

with open(markdown_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_table = False
table_data = []

def make_table(data, document):
    if not data or len(data) < 3: return
    headers = [col.strip() for col in data[0].split('|') if col.strip()]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h.replace('*', '')
        
    for row_line in data[2:]:
        row_cols = [col.strip() for col in row_line.split('|') if col.strip()]
        row_cells = table.add_row().cells
        for i, c in enumerate(row_cols):
            if i < len(row_cells):
                row_cells[i].text = c.replace('*', '')

for line in lines:
    line = line.strip()
    
    if line.startswith('|'):
        in_table = True
        table_data.append(line)
        continue
    elif in_table:
        make_table(table_data, doc)
        in_table = False
        table_data = []
    
    if not line or line.startswith('---'):
        continue
        
    # Headers
    if line.startswith('# '):
        doc.add_heading(line[2:].strip().replace('*', ''), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip().replace('*', ''), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip().replace('*', ''), level=3)
    # Lists
    elif line.startswith('* '):
        doc.add_paragraph(line[2:].strip().replace('*', ''), style='List Bullet')
    elif line[0].isdigit() and line[1:3] == '. ':
        doc.add_paragraph(line[3:].strip().replace('*', ''), style='List Number')
    elif line.startswith('>'):
        doc.add_paragraph(line[1:].strip().replace('*', ''), style='Quote')    
    else:
        # Paragraph
        cln = line.replace('**', '')
        doc.add_paragraph(cln)

if in_table:
    make_table(table_data, doc)

output_path = r'D:\nextjs\c2c-book-hub\Final_Pre_Defence_Report.docx'
doc.save(output_path)
print("Successfully created " + output_path)
